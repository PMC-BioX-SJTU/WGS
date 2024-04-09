#coding=utf-8

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import ConfigParser
import string
import os
import argparse
import sys
from PreProcess import *
cf = ConfigParser.ConfigParser()

parser = argparse.ArgumentParser(description='the Step for Mapping and Mark Duplication reads!')
parser.add_argument('-c', required=True, dest='config',action='store',help='config file for the module')
parser.add_argument('-s', required=True, dest='sample',action='store',help='sample info')
parser.add_argument('-o', required=True, dest='outpath4proj', action='store',help='outfile for make')

if len(sys.argv) <= 3: 
	parser.print_help() 
	sys.exit(1) 
else: 
	args = parser.parse_args()
config = args.config
outpath4proj = args.outpath4proj
sample = args.sample
(name,fq)=sample.split('|')

face=u'\n\n\n\n医学外显子测序分析报告\nV1.0\n\n\n\n\n'
info=u'客户单位：______________________\n报告单位：______________________\n联系人：________________________\n联系电话：______________________\n传真：__________________________\n报告日期：______________________\n项目负责人：____________________\n审核人：________________________\n\n'                      
       
h1=u'1 工作流程' 
h1_1=u'1.1 实验流程' 
p1_1_1=u'外显子捕获测序技术是利用探针捕获试剂盒将基因组外显子区域DNA捕获并富集后，进行高通量测序的方法。首先针对基因组上的目标区域设计特异性探针，然后通过液相芯片杂交的方式将基因组上的DNA目标区域捕获下来，再用高通量测序仪进行测序。具体建库流程如下图：'
#fig1='/opt/NfsDir/UserDir/sunfl/Codes/Pipelines/Modules/exon_exp.png'
fig1_anno=u'图1. 液相芯片杂交捕获平台'



h1_1_1=u'1.1.1 DNA质量检测' 
p1_1_1_1=u'在实验的步骤中对DNA样本的检测方法主要有2种：\n\
（1）琼脂糖凝胶电泳分析DNA降解程度，判断样本是否存在污染；\n\
（2）用Nanodrop检测DNA的纯度（OD260/280比值）和浓度。' 
h1_1_2=u'1.1.2文库构建' 
p1_1_2_1=[u'利用Covaris 超声波破碎方法将基因组DNA 随机打断成180-250bp左右，连上接头,通过PCR扩增构建捕获文库；为了达到等量捕获，计算每个样本的浓度；',u'将扩增后的文库样本和生物素标记的探针进行杂交；',u'利用链霉亲和素修饰的纳米磁珠吸附上一步得到的复合体，清洗磁珠，去除未结合序列，然后将捕获的外显子序列洗脱下来，通过PCR扩增构建外显子文库；',u'通过qRT-PCR来估计目标区域的相对富集倍数，判断样本到达指控的要求后，进行上机测序前的准备；',u'将指控合格的文库样本进行上机测序。']

h1_2=u'1.2 信息分析流程图' 
p1_2_1=u'测序产生的数据经过质量控制后，进入到信息分析阶段。通过Illumina测序仪得到的测序数据称为raw reads。这些raw reads中可能还有adaptors、低质量以及含“N”较多的序列，这些reads 被称为contaminated reads。我们去除这些reads，得到最终的clean reads，再通过BWA[1]比对和samtools处理，得到bam文件，用picard去除重复，最终通过比对的结果检测SNP、InDel变异，并对结果进行注释。信息分析流程图如图2所示：'
#fig2='/opt/NfsDir/UserDir/sunfl/Codes/Pipelines/Modules/exon_pipe.jpg'

fig2_anno=u'图2. 医学外显子测序信息分析流程及质控图'

p1_2_2=u'总的来说整个分析过程可分为两大部分：\n\
1.标准分析：\n\
（1）标准分析流程：主要分为Mapping+cleanup、Variant Discovery和Variant Annotation 三步。\n\
（2）质量评估：与标准分析流程匹配的是数据的质量评估。\n\
    高通量测序从样本收集、建库、产生下机数据一直到生物信息分析产生结果步骤众多，每一步的误差都有可能影响最终的实验结果，因此获得高质量的数据是科研及临床应用服务最基本的要求。本流程的质量控制部分主要包含FASTQ下机数据的评估、Mapping+cleanup后针对bam的评估及针对变异位点的评估三部分。符合标准的数据进入到后续的分析流程。\n\
2.个性化分析：\n\
    个性化分析主要是针对每个项目的不同实验设计而言的。比较通用的包含针对突变位点分析其对遗传密码子的影响，并进行基因功能注释，同时标注突变所在的区域等。\n\
另外，对实验样本进行群体分析。根据人群体样本的实际情况，对群体结构、连锁不平衡等群体特征进行分析，重点是关注的表型或致病风险与SNP等遗传标记的关联分析，找到关联显著的位点，并进行相关的统计、注释等。'

##----------------------------------------------##
h2=u'2 生物信息学分析结果' 
h2_1=u'2.1 原始数据格式说明'
p2_1=u'高通量测序仪(如Illumina HiSeq2000/HiSeq2500/Miseq)得到的原始图像数据文件经CASAVA碱基识别(Base Calling)分析转化为原始测序序列（Sequenced Reads），我们称之为 Raw Data或Raw Reads，结果以 FASTQ (简称为fq)文件格式存储，其中包含测序序列（reads）的序列信息以及其对应的测序质量信息。FASTQ文件格式如下：\n\
@HWI-ST1106:749:H0V1GADXX:2:1101:1213:2133 1:N:0:CTTGTA\n\
CGTTGTCAATGGTGTCAGAGGACTCCACCTCCAGGGTAATGGTCTTACCAG\n\
+\n\
@@@DDDDDF8CBDFFGIEHIFII3CFEIIIIIIGBFFFBFEFFICFEFIII0??BFF-=CD@FFFFBAEFFDBDDDDCBACCBBBBBBBB9\n\
1)第一行以“@”开头，随后为Illumina 测序标识别符(Sequence Identifiers)和描述文字(选择性部分)；\n\
2)第二行是碱基序列；\n\
3)第三行以"+"开头，随后为Illumina 测序标识别符(选择性部分)；\n\
4)第四行是对应碱基的测序质量，该行中每个字符对应的 ASCII 值减去 33，即为对应第二行碱基的测序质量值。\n\
5)通过使用第四行中每个字符对应的ASCII值进行计算，即得到对应第二行碱基的测序质量值。如果测序错误率用e表示，Illumina HiSeq 2500的碱基质量值用Qphred表示，则有下列关系：' 
#fig_s1='/opt/NfsDir/UserDir/sunfl/Codes/Pipelines/Modules/quanlity_calu.png'
p2_1add=u'每个碱基测序错误率是通过测序 Phred 数值（Phred score，Qphred）通过公式一转化得到，而Phred 数值是在碱基识别（Base Calling）过程通过一种预测碱基判别发生错误概率模型计算得到的，对应关系如下表所显示：'

table1_tit=u'表1 Phred 数值与碱基识别错误率的对应关系'
table1=[[u'Phred 分值',u'不正确的碱基识别',u'碱基正确识别率','Q-score'],['10','1/10','90%','Q10'],['20','1/100','99%','Q20'],['30','1/1000','99.9%','Q30'],['40','1/10000','99.99%','Q40' ]]
##--------------------------------------##
h2_2=u'2.2 原始数据的质控及统计'
h2_2_1=u'2.2.1 碱基测序的质量分布' 
p2_2_1=u'在Hiseq2500 测序系统测序时，首先会对文库进行芯片制备，目的是将文库DNA模板固定到芯片上，在固定DNA 模板的过程中，每个DNA 分子会形成一个簇，一个簇就是一个测序位点，在进行固定过程中极少量的簇与簇之间物理位置会发生重叠，在测序时，测序软件通过前4个碱基对这些重叠的点进行分析和识别，将这些重叠点位置分开，保证每个点测到的是一个DNA 分子，因此测序序列5′端前几个碱基的错误率相对较高。另外测序错误率会随着测序序列（Sequenced Reads）的长度的增加而升高，这是由于测序过程中化学试剂的消耗而导致的。因此每个reads 的前几个碱基和后十几个碱基的测序错误率会高于中间的碱基。通过测序软件对每个reads上每个碱基的打分，可绘制每个reads的碱基测序平均质量分布图，所有样本的测序碱基质量分布图的文件目录是QualityControlFq，其中样本'+name+u'的结果如下图所示:'
#fig3_1='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/QualityControlFq/1982.QC_disR1.png'
#fig3_2='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/QualityControlFq/1982.QC_disR2.png'
fig3_anno=u'图3. 样本的测序碱基质量值分布\n\
注：横坐标为reads的碱基位置，纵坐标为所有reads在该位置上的质量分布。上部分为双端测序序列的第一端测序reads的碱基质量值分布情况，下半部分为另一端测序reads的碱基质量值分布情况。'


h2_2_2=u'2.2.2 碱基类型分布'
p2_2_2=u'碱基类型分布检查可用于检测样本汇总有无AT、GC的分离现象，而这种现象来源于建库或测序等环节，且会影响到后续的分析。高通量测序是基因组随机打断成的DNA片段，由于其位点在基因组上的分布是近似均匀的，因此，G/C、A/T含量也是近似均匀的。根据切比雪夫大数定律，可知在每个测序环节上，GC、AT含量应当分别相等，且等于基因组的GC、AT含量。同样因为重叠簇的关系会导致样品前几个碱基AT、GC不等波动较大，高于其它的测序区段，其它区段的GC、AT含量相等，且分布均匀无分离现象，所有样本的测序碱基类型分布图的文件目录是QualityControlFq，其中样本'+name+u'的结果如下图所示：'

#fig4_1='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/QualityControlFq/1982.Nu_distR1.png'
#fig4_2='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/QualityControlFq/1982.Nu_distR2.png'
fig4_anno=u'图4. 样本'+name+u'的碱基类型分布\n\
注：横坐标为reads的碱基位置，纵坐标为所有reads在该位置上的碱基比例。上部分为双端测序序列的第一端测序reads的碱基比例情况，下半部分为另一端测序reads的碱基比例情况。其中碱基N为粉色，碱基T为黄色，碱基G为绿色，碱基C为红色，碱基A为蓝色。'

##=------------------------------------------##
h2_2_3=u'2.2.3 测序数据的过滤'
p2_2_3=u'测序完成下机后得到的原始数据，并不都是有效的，里面含有带接头的，重复的，测序质量低的reads，这些reads会影响后续分析，我们必须对下机的原始数据进行过滤，得到干净有效的reads。\n\
数据的过滤步骤：\n\
    1. 去除含adaptor的reads\n\
    2. 去除N的比例大于10%的reads\n\
    3. 去除低质量reads（质量值Q <= 3的碱基数占整个reads的50％以上）\n\
    4. 获得 Clean reads \n\
数据过滤后，数据过滤后，各样本产出数据评估结果如下表所示'
table2=[]
table2_tit=u'表2 样品测序数据评估统计'
table2_anno=u'Sample ID:样本名 \n\
TOTAL_READS: The total number of reads in the input file(所有reads数目) \n\
READ_LENGTH: The average read length of all the reads (所有reads的平均长度) \n\
TOTAL_BASES: The total number of bases in all reads(所有reads的碱基数目)\n\
Q20_BASES:The number of bases in all reads that achieve quality score 20 or higher(所有reads中质量分数大于20的碱基数目)\n\
Q20_BASES_RATE: The percentage of bases in all reads that achieve quality score 20 or higher(所有reads中质量分数大于20的碱基数目的百分比)\n\
Q30_BASES: The number of bases in all reads that achieve quality score 30 or higher(所有reads中质量分数大于30的碱基数目)\n\
Q30_BASES_RATE: The percentage of bases in all reads that achieve quality score 30 or higher(所有reads中质量分数大于30的碱基数目的百分比)\n'


h2_3=u'2.3 Clean reads和参考基因组比对与外显子捕获统计' 
h2_3_1=u'2.3.1 测序数据和参考基因组的比对结果质控及统计'
p2_3_1=u'样本测序数据经过质控筛选后，和参考基因组序列进行比对。在此我们用BWA工具将样本和人类参考基因组序列进行比对。参考基因组是组装的DNA，可能来源于多个个体的DNA序列，与任意个体的DNA序列均不完全一致。不同的数据库收录了不同版本的人类参考基因组，比如hg19、GRCh37、ensembl75等，这些是分别由UCSC、NCBI、EMBLE颁布的参考基因组。hg系列是UCSC颁布的人类参考基因组，也是目前使用频率最高的参考基因组。在本WES分析流程中，我们采用hg19作为参考基因组，hg19对应于GRCh37，基因组序列相同，是在GRCh36（hg18）的基础上，更新了人类基因组的所有染色体，接合了25处序列的gaps，修正超过了150处有问题的序列，并且加了9个alternate loci后得到的版本。\n\
由于有效数据量及数据比对情况对分析准确性有很大影响，所以在突变检测分析前，需要进行数据产出量统计、比对结果、以及对数据进行适当筛选，以有利于后续分析及结果的可靠性。\n\
将测序数据与参考基因组进行比对分析，并除去结果文件中的冗余序列，可对数据测序覆盖区域、测序数据覆盖深度等做出综合评价。\n\
在得到clean reads后，我们采用BWA软件将clean reads与参考基因组hg19做比对，初始比对结果为sam格式，再利用samtools软件将结果转为bam格式并排序。如果一个样品的结果包含多个文库，则用samtools将多个文库的bam结果合并。然后用picard标注并去除重复序列，然后进行数据基本信息统计及map比对统计，详细结果如下：'
table3=[]
table3_tit=u'表3 样品比对结果评估统计'

table3_anno=u'Sample ID	样本名\n\
TOTAL_READS: The total number of reads including all PF and non-PF reads(测序数据中包括质控过滤掉的和未过滤掉的所有的reads数)\n\
PF_READS: The number of PF reads where PF is defined as passing Illumina\'s filter(原始测序数据经过质控过滤后得到的reads数)\n\
PCT_PF_READS: The percentage of reads that are PF (PF_READS / TOTAL_READS)质控过滤后得到的reads数占总reads数的比例\n\
PCT_PF_READS_ALIGNED: The percentage of PF reads that aligned to the reference sequence. PF_READS_ALIGNED / PF_READS(质控过滤后得到的reads中比对到参考基因组上的reads数)\n\
PF_HQ_ALIGNED_READS: The number of PF reads that were aligned to the reference sequence with a mapping quality of Q20(比对的质量值到达Q20的PF的reads数)\n\
PF_MISMATCH_RATE: The rate of bases mismatching the reference for all bases aligned to the reference sequence.(Reads中错配的碱基占所有碱基的比例)\n\
PF_HQ_ERROR_RATE: The percentage of bases that mismatch the reference in PF HQ aligned reads(在PF HQ比对的reads中与参考序列错配的碱基的百分比)\n\
PF_INDEL_RATE: The number of insertion and deletion events per 100 aligned bases. Uses the number of events as the numerator, not the number of inserted or deleted bases.(每100个比对碱基插入和缺失事件的数目，使用事件的数目而不是插入或者缺失碱基的数目作为计数器。)\n\
PCT_READS_ALIGNED_IN_PAIRS: The percentage of reads whose mate pair was also aligned to the reference. READS_ALIGNED_IN_PAIRS / PF_READS_ALIGNED (其互补链能比对到参考序列的PF reads数量占比对到参考基因组的总的PF reads的百分比)\n\
STRAND_BALANCE:The number of PF reads aligned to the positive strand of the genome divided by the number of PF reads aligned to the genome.(比对道基因组正义链的PF reads占比对到基因组上总的reads的比例)\n\
PCT_CHIMERAS: The percentage of reads that map outside of a maximum insert size (usually 100kb) or that have the two ends mapping to different chromosomes.(定位到一个最大插入片段（通量100kb）或者那些有两个端点定位到到不同的染色体上的reads的百分比)\n'


h2_3_2=u'2.3.2 样本交叉污染质控' 
p2_3_2=u'在比对得到结果进行下一步分析之前，需要确认样本产生的数据是否符合预期，即判断不同样本之间是否存在着交叉污染。在此可以通过VerifyBAMID软件来对样本的测序数据进行检测，以判断该测序数据是来源于样本自身或是来源于其它的样本。VerifyBAMID软件的运行，以千人基因组中的亚洲人群的.vcf文件和经过标记重复及BQSR校正后的.BAM文件为输入文件。每个单一样本运行产生对应的6个对应的文件。根据Broad研究所的标准，当FREEMIX值大于0.075时，认为样本被污染。具体分析结果如下：'
table4=[]
table4_tit=u'表4 样品交叉污染检测评估统计'
table4_anno=u'Sample_ID: 样本名称\n\
#SNPs:# of SNPs passing the criteria from the VCF file(来源于VCF文件的通过标准的SNPs数)\n\
#READS: Total # of reads loaded from the BAM file(来源于BAM文件的reads数)\n\
AVG_DP: Average sequencing depth at the sites in the VCF file(位点的平均测序深度)\n\
FREEMIX: Sequence-only estimate of contamination (0-1 scale)(在仅有测序数据模式下检测样本交叉污染的估计值)'

h2_3_3=u'2.3.3 外显子测序探针捕获质控' 
p2_3_3=u'对外显子及临近侧翼区域的比对结果进行统计，并依据其覆盖深度和比例来判断外显子捕获的效率。样品各区域的比对结果统计如下（由于样品量多时，不便于展示，所以我们只展示其中一个样品的统计表）：'
table5=[]
table5_tit=u'表5 样品外显子捕获统计'


table5_anno=u'Sample ID	样本名\n\
TOTAL_READS:The total number of reads in the SAM or BAM file examined(总的reads数)\n\
PF_READS: The total number of reads that pass the vendor s filter.(通过vendor标准过滤后的reads总数)\n\
PF_UNIQUE_READS: The number of PF reads that are not marked as duplicates(经过过滤了的reads中，没有被标记为重复的reads数)\n\
PCT_PF_UQ_READS: The fraction of PF_UNIQUE_READS from the TOTAL_READS, PF_UNIQUE_READS/TOTAL_READS(经过过滤了的reads中，没有被标记为重复的reads数占总reads数的比例)\n\
ON_TARGET_BASES: The number of PF_BASES_ALIGNED that are mapped to the targeted regions of the genome(经过过滤的，没有被标记为重复的reads中比对到参考基因组上诱饵区域内的base数)\n\
NEAR_TARGET_BASES: The number of PF_BASES_ALIGNED that are mapped to within a fixed interval containing a targeted region, but not within the targeted section per se.(经过过滤的，没有被标记为重复的reads中比对到参考基因组上包含靶向捕获区域的固定间隔内，但不在靶向捕获区段本身内（诱饵区域附近）的base数)\n\
OFF_BAIT_BASES: The number of PF_BASES_ALIGNED that are mapped away from any baited region.(经过过滤的，没有被标记为重复的reads中比对到参考基因组上远离诱饵区域内的base数)\n\
PCT_SELECTED_BASES: The fraction of PF_BASES_ALIGNED located on or near a baited region (ON_BAIT_BASES + NEAR_BAIT_BASES)/PF_BASES_ALIGNED.(经过过滤的，没有被标记为重复的reads中，比对到参考基因组上诱饵区域内及其附近的base数占这部分reads中比对到参考基因组上的总的base数的比例)\n\
PCT_OFF_BAIT: The fraction of PF_BASES_ALIGNED that are mapped away from any baited region, OFF_BAIT_BASES/PF_BASES_ALIGNED(经过过滤的，没有被标记为重复的reads中，比对到参考基因组上非诱饵区域内的base数占这部分reads中比对到参考基因组上的总的base数的比例)\n\
MEAN_TARGET_COVERAGE: The mean coverage of a target region.(靶标区间的均值)\n\
MEDIAN_TARGET_COVERAGE: The median coverage of a target region.(靶标区间的中位数)\n\
FOLD_ENRICHMENT: The fold by which the baited region has been amplified above genomic background.(在参考基因组背景上捕获下的bait区域被扩增的倍数)\n\
ZERO_CVG_TARGETS_PCT: The fraction of targets that did not reach coverage=1 over any base.(所有碱基中不能达coverage=1的 targets 的比例)\n\
PCT_EXC_DUPE:The fraction of aligned bases that were filtered out because they were in reads marked as duplicates.(过滤掉的比对碱基的比例，因为这些碱基在reads中被标记为duplicates)\n\
PCT_EXC_MAPQ: The fraction of aligned bases that were filtered out because they were in reads with low mapping quality.(过滤掉的比对碱基的比例，因为这些碱基在reads中有很低的mapping值)\n\
PCT_EXC_BASEQ: The fraction of aligned bases that were filtered out because they were of low base quality(过滤掉的比对碱基的比例，因为这些碱基有很低的碱基质量)\n\
PCT_EXC_OVERLAP: The fraction of aligned bases that were filtered out because they were the second observation from an insert with overlapping reads(过滤掉的比对碱基的比例，因为他们是来自在重叠reads中插入的再次观察)\n\
PCT_EXC_OFF_TARGET: The fraction of aligned bases that were filtered out because they did not align over a target base(过滤掉的比对碱基的比例，因为他们没有比对上任何一个target碱基)\n\
PCT_TARGET_BASES_1X:The fraction of all target bases achieving 1X or greater coverage(所有目标碱基覆盖度1X或更高的比例)\n\
PCT_TARGET_BASES_2X:The fraction of all target bases achieving 2X or greater coverage(所有目标碱基覆盖度2X或更高的比例)\n\
PCT_TARGET_BASES_10X:The fraction of all target bases achieving 10X or greater coverage(所有目标碱基覆盖度10X或更高的比例)\n\
PCT_TARGET_BASES_20X:The fraction of all target bases achieving 20X or greater coverage(所有目标碱基覆盖度20X或更高的比例)\n\
PCT_TARGET_BASES_30X:The fraction of all target bases achieving 30X or greater coverage(所有目标碱基覆盖度30X或更高的比例)\n\
PCT_TARGET_BASES_40X:The fraction of all target bases achieving 40X or greater coverage(所有目标碱基覆盖度40X或更高的比例)\n\
PCT_TARGET_BASES_50X:The fraction of all target bases achieving 50X or greater coverage(所有目标碱基覆盖度50X或更高的比例)\n\
PCT_TARGET_BASES_100X:The fraction of all target bases achieving 100X or greater coverage(所有目标碱基覆盖度100X或更高的比例)\n'

p2_3_4=u'该样本的插入片段长度分布图如下所示，其中虚线为累计掺入片段长度分布曲线。'


#fig5='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/QualityControlBam/1982.insert_size.png'
fig5_anno=u'图5. 样本'+name+u'所有reads的插入片段直方图\n\
注：横坐标为插入片段的长度，纵坐标为不同长度插入片段的数量统计。'

p2_3_5=u'该样本的测序深度及累计测序深度分布图如下所示'

#fig6='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/QualityControlBam/1982_histPlot.png'
fig6_anno=u'图6. 样本的测序深度分布图\n\
注：横坐标表示测序深度，纵坐标为百分比。'

#fig7='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/QualityControlBam/1982_cumuPlot.png'
fig7_anno=u'图7. 样本累计测序深度分布图\n\
注：横坐标为测序深度，纵坐标为大于测序深度的Read所占的比例。'

h2_4=u'2.4 突变检测及注释'
h2_4_1=u'2.4.1 样本与参考基因组间的SNP，INDEL检测'
p2_4_1=u'测序产生的reads首先与基因组参考序列进行比对，其次使用picard、GATK[2]等相关工具进行SNP和INDEL的检测、过滤，得到变异位点集，最后采用ANNOVAR软件[3]对所有位点进行注释。详细流程如下：\n\
（1） 对于BWA 比对得到的结果，使用Picard 的Mark Duplicate 工具标记重复序列，屏蔽PCR-duplication 的影响。\n\
（2） 使用GATK 进行INDEL Realignment，即对存在插入缺失比对结果附近的位点进行局部重新比对，校正由于插入缺失引起的比对结果错误。\n\
（3） 使用GATK 进行碱基质量值再校准（Base Recalibration），对碱基的质量值进行校正。\n\
（4） 使用GATK 进行变异检测（variant calling），主要包括SNP 和INDEL。\n\
（5） 使用GATK 对得到的变异结果进行校正，选取可靠的变异结果。\n\
    具体文档可参考GATK 官方网站的Best Practice：（https://www.broadinstitute.org/gatk/guide/best-practices?bpm=DNAseq#variant-discovery-ovw）以上过程中（1）-（4）完成了变异位点的查找，生成一个包含raw variants的VCF文件，我们根据GATK官方网站的Best Practice提供的标准对样本实行严格的质控过滤，以保证变异结果的可靠性。根据实验样本数的多少，GATK官方网站的Best Practice提供了两种质控过滤的标准少数样本的hard filtering和多个样本的VQSR。变异结果我们经过hard filtering严格的过滤，保证变异结果的可靠性。过滤时采用GATK中的过滤参数及阈值，符合如下标准的变异位点留下:'

table6_tit=u'表6 样品初始变异过滤校正标准'
table6=[['Sample ID',u'样本名','SNP','INDEL'],['LOW_READ_DEPTH',u'低读取深度','<10',''],['Quality Depth',u'由非参考样品的未过滤的深度分开的变异置信度','<2',''],['FisherStrand',u'使用Fisher’s Exact Test检测reads中链的bias的Phred-scaled p值。更多的bias只是为假阳性calls','>60','>200'],['MQRankSum',	u'这是对定位质量（参考碱基的reads vs. alternate allele）来说来自Mann-Whitney秩和检验的基于u的z近似值。注意定位质量秩和检验在没有显示出参考基因和alternate allele的混合reads下不能被计算出位点，例如这仅仅应用于杂合calls。','<-12.5',''],['RMSMappingQuality',u'所有样品中reads的定位质量的Root Mean Square','<40',''],['ReadPosRankSumTest',u'这是对从带有alternate allele的reads的末端的距离的Mann-Whitney 秩和检验的基于u的z的近似值。如果可选等位基因仅仅在临近reads末端被看见，这样就显示错误。注意，reads位置的秩和检验在没有显示出参考基因和alternate alleles的混合reads下不能被计算出位点，例如这仅仅应用于杂合calls。',u'<-0.8','<-20']]



p2_4_1_add=u'在经过上述质控筛选步骤后，生成的变异结果使用vcf 文件格式展示。vcf 文件的详细说明信息见网页：http://gatkforums.broadinstitute.org/discussion/1268/how-should-i-interpret-vcf-files-produced-by-the-gatk。vcf 文件包括注释行、标题行和数据行三部分。其中注释行包含文件数据行的INFO 和FORMAT 列中使用的各种标识符的意义解释，而标题行和数据行包含各样品的变异检测结果信息，标题行中各数据项对应意义如下：'
table7_tit=u'表7 变异结果vcf文件各指标意义'
table7=[[u'标题',u'示例',u'意义说明'],['CHROM','Chr1',u'参考序列的染色体名称'],['POS','5634',u'参考序列位点坐标'],['ID','.',u'标识符'],['REF',	'G',u'参考序列对应位置碱基'],['ALT','A','SNV'],['QUAL','140.84',u'测序质量值'],['FILTER','PASS',u'过滤状态'],['INFO','ANNOTATIONS',u'附加信息'],['FORMAT','GT:AD:DP:GQ:PL',u'基因型信息格式'],['R01', '1/1:0','6:6:18:169','18']]


h2_4_2=u'2.4.2 SNP和INDEL结果的注释数据库简介'
p2_4_2=u'在获得最终的变异位点之后，我们用ANNOVAR对变异结果进行注释ANNOVAR[3]是一款当前比较常用且全面的注释变异（SNP、Small INDEL）位点的软件，广泛的用于对人类变异结果的注释，支持大多 UCSC 上现有的数据库，以及用户自定义的数据库。根据变异位点在参考基因组上的位置以及参考基因组上的基因位置信息，可以得到变异位点在基因组发生的区域（基因间区、基因区或UTR 区等），以及变异产生的影响（同义/非同义突变等）。软件使用vcf 格式文件进行输入，具体说明可参见 ANNOVAR 的说明文档：http://annovar.openbioinformatics.org/en/latest/\n\
数据库注释：'
p2_4_2_add=u'1) RefSeq genes，GENCODE genes annotation\n\
2) Allele frequency in the 1000 Genome Project, NHLBI-ESP 6500\n\
3) GwasCatalog：已经发表的各种疾病的GWAS结果（NHGRI\'s collection ofGenome-Wide Association Studies SNVs）[6]\n\
4) dbSNP138\n\
5) LJB26 (SIFT ，PolyPhen 2，LRT，MutationAssessor，GERP++，PhyloP，SiPhy等注释)。该数据库对变异位点进行SIFT（Sorting Intolerant From Tolerant） [4] 和PolyPhen-2 （Polymorphism Phenotyping v2） [5]预测；SIFT是根据氨基酸替换保守性对蛋白质功能产生影响的预测软件，它可以判断出这个氨基酸置换在蛋白质功能上是无害的（functionally neutral）的还是有害的（deleterious），一般认为<=0.05为有害的突变；PolyPhen-2也是一种预测氨基酸置换对蛋白质结构和功能影响的工具，一般认为>= 0.909为有害突变。\n\
6) miRBase 注释（snoRNA and microRNAs）\n\
7) genomicSuperDups：基因组中的重复片段\n\
8) phastConsElements46way：通过phastCons对脊椎动物的全基因组比对生成的保守区域，根据用于比对的物种数目，分为17way, 28way, 30way, 44way等。\n\
9) OMIM（Online Mendelian Inheritance in Man）：http://www.omim.org/。 该数据库中包括了所有的15000多种已知的单基因疾病表型和其致病基因间的关系。\n\
10) dbnsfp30a：非同义突变注释（db non-synonymous variants annotation）由于历史原因，在ANNOVAR中被命名为ljb而不是dbNSFP。截止目前，最新的ljb数据库为dbnsfp30a。\n\
11) 1000g2015aug_all：最新的千人基因组数据库，包含6个人种的基因突变频率，AFR (非洲)，AMR (混合美洲)，EAS (东亚)， EUR (欧洲)， SAS (南亚)\n\
12) gerp ++：基于人类基因组选择性约束的90亿突变的功能预测分数。 您可以选择使用gerp ++ gt2，因为它只包括大于2的RS分数，这提供了高灵敏度。\n\
13) clinvar_20160302：每个变体的具有单独列（CLINSIG CLNDBN CLNACC CLNDSDB CLNDSDBID）的ClinVar数据库。它存储有关变异和人类健康之间关系的信息。\n\
14) cosmic70：最新的COSMIC数据库，包含癌症的体细胞突变和癌症的每个亚型发生的频率。\n\
15) exac03：最新的ExAC数据库中的等位基因频率，包含所有人的，AFR（非洲），AMR（混合美洲），EAS（东亚），FIN（芬兰），NFE（非芬兰语），OTH （南亚）和其他人种。\n\
16) esp6500siv2：最新的NHLBI-ESP项目，包含6500个外显子数据。 三个独立的关键词用于3个人口分组：esp6500siv2_all，esp6500siv2_ea，esp6500siv2_aa。\n\
17) avSNP147：dbSNP 147的缩写版本。\n\
18) popfreq_max_20150413：popfreq_max数据库包含来自几个数据库（包含种群频率）的最大等位基因频率，包括千人基因组（ALL + 5种族群），ESP6500（ALL + 2种族群），ExAC（ALL + 7种族群），CG46。 popfreq_all数据库包含多个列，代表来自这些种群频率数据库的所有等位基因频率。20150413是数据发布日期。\n\
19) CADD（Combined AnnotationDependent Depletion）是基于SVM对多个分数进行综合的分数。它为人类基因组中的每个可能的突变分配一个分数，因此可以评估非编码突变以及编码突变。'

h2_4_3=u'2.4.3 SNP和INDEL位点注释统计'
p2_4_3_add=u'在经过以上步骤过滤筛选后得到了较为可靠的variants的集合，为了进一步提高所得变异位点的可靠性，最大化真实的位点数量、降低假阳性，在将得到的变异位点做注释之前，可将上述初步过滤后的变异位点集合和已知的标准品序列进行比较，完成进一步的质控过滤。判断出这些变异位点是否和已知的标准品突变位点重合，理论上不同样本间和标准品的重合数量大致一致。经过此步过滤可以筛选查找出离群的样本。由于现实情况中标准品的制备不易，我们采取1000G和dbSNP两个公共数据库数据作为标准品，来对样本进行质控校正。\n\
对SNP文件进行注释后，去掉有缺失的SNP位点后，再进行统计和过滤；通过ANNOVAR对SNP文件进行注释后，去掉有缺失的SNP位点后，再进行统计和过滤；本项目各样品的SNP 结果具体统计结果如下表：'
table8_tit=u'表8 VCF中变异位点的注释质控统计结果'
table8=[]

table8_anno=u'Total Variants：检测出的变异位点总数\n\
Het：杂合变异位点数目\n\
Homo：纯合变异位点数目\n\
dbSNP：注释到dbSNP的变异位点数\n\
1000G：注释到千人基因变异位点数\n\
esp6500：注释到esp6500的变异位点数目\n\
ExAC：注释到ExAC的变异位点数目 \n\
Reference_all：注释到dbSNP、1000G、esp6500或ExAC的变异位点数目\n\
Novel：没有在以上数据库发现的变异位点数目\n\
exonic：注释到exonic区的变异位点数目\n\
splicing：注释到splicing区的变异位点数目\n\
exonic;splicing：同时注释到exonic;splicing区的变异位点数目\n\
intergenic：注释到基因间区的变异位点数目\n\
intronic：注释到intronic区的变异位点数目\n\
UTR3：注释到UTR3区的变异位点数目\n\
UTR5：注释到UTR5区的变异位点数目\n\
UTR5;UTR3：注释到UTR5;UTR3区的变异位点数目\n\
ncRNA_exonic：注释到ncRNA_exonic区的变异位点数目\n\
ncRNA_splicing：注释到splicing区的变异位点数目\n\
ncRNA_exonic,splicing：注释ncRNA_exonic,splicing区的变异位点数目\n\
ncRNA_intronic：注释到ncRNA_intronic区的变异位点数目\n\
ncRNA_UTR5：注释到ncRNA_UTR5区的变异位点数目\n\
ncRNA_UTR3：注释到ncRNA_UTR3区的变异位点数目\n\
upstream：注释到upstream区的变异位点数目\n\
downstream：注释到downstream区的变异位点数目\n\
upstream;downstream：注释到upstream;downstream区的变异位点数目\n\
unknown：注释到unknown区的变异位点数目\n\
synonymous：同义突变变异位点数目\n\
nonsynonymous：非同义突变变异位点数目\n\
frameshift deletion：移码改变的 deletion变异位点数目\n\
frameshift insertion：移码改变的indersiton变异位点数目\n\
nonframeshift deletion：非frameshift deletion变异位点数目\n\
nonframeshift insertion：非frameshift indersiton变异位点数目\n\
stopgain：注释到stopgain区的变异位点数目\n\
stoploss：注释到stoploss区的变异位点数目\n\
Ti/Tv|SNP/INDEL Rate：转换(transition)数/颠换(transversion)的比例|SNP/INDEL变异比例'


p2_4_2_add1=u'提取注释结果文件中的样本中各类型的SNP变异位点数量，作图如下'
#fig8='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/Annotation/1982.SNP_Spectrum.png'
fig8_anno=u'图8. 样本'+name+u'中不同SNP变异位点类型的数量统计\n\
注：横坐标表示SNP的类型，纵坐标表示对应SNP类型的个数。'

p2_4_2_add2=u'根据样品的Clean Reads 在参考基因组UCSC hg19 上的定位结果，检测样品与参考基因组之间是否存在小片段的插入与缺失（Small INDEL）。样品的插入缺失使用GATK 检测。INDEL 变异一般比SNV 变异少，同样反映了样品与参考基因组之间的差异，并且编码区的INDEL 会引起移码突变，导致基因功能上的变化。根据样品检测得到的Small INDEL 位点在参考基因组上的位置信息，对比参考基因组的基因、exonic 位置等信息(一般在gff 文件中)，可以注释INDEL 位点是否发生在基因间区、基因区或exonic 区、是否为移码突变等。通过ANNOVAR对VCF文件中的SNP位点进行注释后，本项目'+name+u'样本的INDEL 结果具体统计结果如下表所示： '
table9_tit=u'表9 VCF中INDEL变异位点的注释质控统计结果'
table9=table1

#fig9='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/Annotation/1982.InDel_Distribution.png'
fig9_anno=u'图9. 样本'+name+u'中不同INDEL变异类型（种类及片段长度）的数量统计\n\
注：INDEL长度分布图：横坐标表示INDEL长度，纵坐标表示对应长度INDEL的个数。'

#fig8='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/Annotation/1982.InDel_Distribution.png'
#fig8_anno=u'图8. 样本'+name+u'中分布在不同区域的变异位点数量统计图\n\
#注：横坐标为分布的不同区域，纵坐标为对应于不同区域中的变异位点的统计数量。'
#fig9='/opt/NfsDir/UserDir/sunfl/Codes/test/1982/Annotation/1982.InDel_Distribution.png'
#fig9_anno=u'图9. 样本'+name+u'中包含的不同变异类型的数量统计图\n\
#注：横坐标为不同的突变类型，纵坐标为对应于不同变异类型的统计数量。'

h3=u'3 所用软件及数据库'
h3_1=u'3.1 软件及工具：'
p3_1='BWA (Burrows-Wheeler Aligner) 0.7.5a-r405\n\
samtools (Sequence Alignment/Map Tools) 0.1.19\n\
picard-tools 1.57\n\
Genome Analysis Toolkit 2.5-2-gf57256b\n\
ANNOVAR'
h3_2=u'3.2 数据库信息：'
p3_2='dbSNP:    ftp://ftp.ncbi.nih.gov/SNP/database/\n\
EBI FTP:    ftp://ftp.1000genomes.ebi.ac.uk/vol1/ftp/\n\
NCBI FTP:    ftp://ftp-trace.ncbi.nih.gov/1000genomes/ftp/\n\
Human reference genome (NCBI build37) at UCSC\n\
hg19:    http://hgdownload.cse.ucsc.edu/goldenPath/hg19/bigZips/\n\
CCDS ftp://ftp.ncbi.nlm.nih.gov/pub/CCDS/current_human/'

h4=u'4 参考文献'
p4=u'1. Pleasance, E.D., et al., A comprehensive catalogue of somatic mutations from a human cancer genome. Nature, 2010. 463(7278): p. 191-6.\n\
2. Coffey, A.J., et al., The GENCODE exome: sequencing the complete human exome. Eur J Hum Genet, 2011. 19(7): p. 827-31.\n\
3. Yamaguchi, T., et al., Exome resequencing combined with linkage analysis identifies novel PTH1R variants in primary failure of tooth eruption in Japanese. J Bone Miner Res, 2011. 26(7): p. 1655-61.\n\
4. Koboldt, D.C., et al., VarScan 2: somatic mutation and copy number alteration discovery in cancer by exome sequencing. Genome Res, 2012. 22(3): p. 568-76.\n\
5. Lee, W., et al., Bi-directional SIFT predicts a subset of activating mutations. PLoS One, 2009. 4(12): p. e8311.\n\
7. Ng, P.C. and S. Henikoff, SIFT: Predicting amino acid changes that affect protein function. Nucleic Acids Res, 2003. 31(13): p. 3812-4.\n\
6. Adzhubei, I.A., et al., A method and server for predicting damaging missense mutations. Nat Methods, 2010. 7(4): p. 248-9.'
h5=u'5 联系我们'
p5=u'地址：北京市海淀区苏州街16号神州数码大厦\n\
电话：010-62693727\n\
邮箱：smgene@yeah.net'



##------------------------------------------------------------------##
def Define_style():
	style = document.styles['Normal']
	font = style.font
	font.name=u'宋体'
	#font.name = 'Times New Roman'
	r = style._element
	r.rPr.rFonts.set(qn('w:eastAsia'), font.name)
	font.size = Pt(10)
	return(style)
def Out_Format1(out, fontname,btype,size,align):
	paragraph = document.add_paragraph()
	run = paragraph.add_run(out)
	font_name=fontname
	run.bold=btype
	run.font.size = Pt(size)
	run.font.name = font_name
	run.font.name = 'Times New Roman'
	r = run._element
	r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = align

def Out_Format2(out, style,align,indent):
	paragraph = document.add_paragraph(out,style=style)
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
	paragraph_format.first_line_indent = Inches(indent)

def Out_Format3(out, align):
	paragraph = document.add_paragraph(out)
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = align

def Out_List(parag,style):
	for i in range(0,len(parag)):
		paragraph = document.add_paragraph(parag[i],style=style)

def Out_Fig1(fig, width,fig_anno,fontname,size,align,indent):
	document.add_picture(fig, width=Inches(width))
	paragraph = document.add_paragraph()
	run = paragraph.add_run(fig_anno)
	font_name=fontname
	run.font.size = Pt(size)
	run.font.name = font_name
	run.font.name = 'Times New Roman'
	r = run._element
	r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
	paragraph_format = paragraph.paragraph_format
	paragraph_format.alignment = align
	paragraph_format.first_line_indent = Inches(indent)

def Out_table(r,c,tab):
	table = document.add_table(rows=r, cols=c,style='Table Grid')
	for i in xrange(r):
		for j in xrange(c):
			table.cell(i,j).text =tab[i][j]
		#	table.cell(i,j).text.font.size=Pt(7) 

def Out_table1(r,c,tab):
        table = document.add_table(rows=r, cols=c,style='Light Shading')
	for i in xrange(r):
		for j in xrange(c):
			table.cell(i,j).text =tab[j][i]

def Read_Tables(table):
	out=[]
	t=open(table,'r')
	lines=t.readlines()
	for i in range(0,len(lines)):
		value=lines[i].strip('\n').split('\t')
		out.append(value)
	return(out)

def Check_Exists(file):
	for i in range(0,len(file)):
		if not os.path.exists(file[i]):
			print "We cannot find the file "+file[i]+"!!"
			sys.exit(1)

if __name__ == '__main__':
	document = Document()
	style=Define_style()
	#requirements preparation
	soft_opt=["module_path"]
	param_opt=["skip"]
	ref_opt=["ref"]
	opt=[soft_opt,param_opt,ref_opt]
	sec=['software','Report4WESBasic.parameter','reference']
	### SampleAndOutpat, software, parameter,reference
	requirements=Input_Requiremens(config,sec,opt,sample,"Report4WESBasic",outpath4proj)
	(tab,fig)=requirements[0][1].split('|')
	
	tables=tab.split(':')
	tables[0]=outpath4proj+'/'+name+'/QualityControlBam/'+name+tables[0]
	tables[1]=outpath4proj+'/'+name+'/QualityControlBam/'+name+tables[1]
	tables[2]=outpath4proj+'/'+name+'/QualityControlBam/'+name+tables[2]
	tables[3]=outpath4proj+'/'+name+'/QualityControlBam/'+name+tables[3]
	tables[4]=outpath4proj+'/'+name+'/Annotation/'+name+tables[4]
	Check_Exists(tables)

	figures=fig.split(':')
	figures[0]=requirements[1][0]+'/'+figures[0]
	figures[1]=requirements[1][0]+'/'+figures[1]
	figures[2]=requirements[1][0]+'/'+figures[2]
	figures[3]=outpath4proj+'/'+name+'/QualityControlFq/'+name+figures[3]
	figures[4]=outpath4proj+'/'+name+'/QualityControlFq/'+name+figures[4]
	figures[5]=outpath4proj+'/'+name+'/QualityControlFq/'+name+figures[5]
	figures[6]=outpath4proj+'/'+name+'/QualityControlFq/'+name+figures[6]
	figures[7]=outpath4proj+'/'+name+'/QualityControlBam/'+name+figures[7]
	figures[8]=outpath4proj+'/'+name+'/QualityControlBam/'+name+figures[8]
	figures[9]=outpath4proj+'/'+name+'/QualityControlBam/'+name+figures[9]
	figures[10]=outpath4proj+'/'+name+'/Annotation/'+name+figures[10]
	figures[11]=outpath4proj+'/'+name+'/Annotation/'+name+figures[11]
	Check_Exists(figures)
	
	### write to word
	Out_Format1(face,u'宋体' ,True,28,WD_ALIGN_PARAGRAPH.CENTER)
	Out_Format1(info,u'宋体' ,True,11,WD_ALIGN_PARAGRAPH.CENTER)
	document.add_page_break()

	document.add_heading(h1,1)
	document.add_heading(h1_1,2)
	Out_Format2(p1_1_1, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Fig1(figures[0], 4,fig1_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)
	
	document.add_heading(h1_1_1,3)
	Out_Format3(p1_1_1_1, WD_ALIGN_PARAGRAPH.LEFT)
	document.add_heading(h1_1_2,3)
	Out_List(p1_1_2_1,'List Bullet')

	document.add_heading(h1_2,1)
	Out_Format2(p1_2_1, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Fig1(figures[1], 4,fig2_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)
	Out_Format3(p1_2_2, WD_ALIGN_PARAGRAPH.LEFT)

	document.add_heading(h2,1)
	document.add_heading(h2_1,2)
	Out_Format2(p2_1, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)

	document.add_picture(figures[2], width=Inches(1.5))
	Out_Format2(p2_1add, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Format1(table1_tit,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.CENTER)
	Out_table(5,3,table1)

	document.add_heading(h2_2,2)
	document.add_heading(h2_2_1,3)
	Out_Format2(p2_2_1, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	document.add_picture(figures[3], width=Inches(4))
	Out_Fig1(figures[4], 4,fig3_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)

	document.add_heading(h2_2_2,3)
	Out_Format2(p2_2_2, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	document.add_picture(figures[5], width=Inches(4))
	Out_Fig1(figures[6], 4,fig4_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)

	document.add_heading(h2_2_3,3)
	Out_Format2(p2_2_3, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Format1(table2_tit,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.CENTER)
	table2=Read_Tables(tables[0]) ###
	Out_table(8,2,table2)
	Out_Format1(table2_anno,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.LEFT)

	document.add_heading(h2_3,2)
	document.add_heading(h2_3_1,3)
	Out_Format2(p2_3_1, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Format1(table3_tit,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.CENTER)
	table3=Read_Tables(tables[1])
	Out_table(12,2,table3)
	Out_Format1(table3_anno,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.LEFT)
	
	document.add_heading(h2_3_2,3)
	Out_Format2(p2_3_2, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Format1(table4_tit,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.CENTER)
	table4=Read_Tables(tables[2])
	Out_table(5,2,table4)
	Out_Format1(table4_anno,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.LEFT)

	document.add_heading(h2_3_3,3)
	Out_Format2(p2_3_3, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Format1(table5_tit,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.CENTER)
	table5=Read_Tables(tables[3])
	Out_table(27,2,table5)
	Out_Format1(table5_anno,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.LEFT)
	Out_Format2(p2_3_4, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Fig1(figures[7], 4,fig5_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)
	Out_Format2(p2_3_5, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Fig1(figures[8], 4,fig6_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)
	Out_Fig1(figures[9], 4,fig7_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)


	document.add_heading(h2_4,2)
	document.add_heading(h2_4_1,3)
	Out_Format2(p2_4_1, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Format1(table6_tit,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.CENTER)
	Out_table(7,4,table6)

	Out_Format2(p2_4_1_add, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Format1(table7_tit,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.CENTER)
	Out_table(11,3,table7)

	document.add_heading(h2_4_2,3)
	Out_Format2(p2_4_2, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Format3(p2_4_2_add, WD_ALIGN_PARAGRAPH.LEFT)
	document.add_heading(h2_4_3,3)
	Out_Format2(p2_4_3_add, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Format2(p2_4_2_add2, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Format1(table8_tit,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.CENTER)
	table8=Read_Tables(tables[4])
	Out_table(37,3,table8)
	Out_Format1(table8_anno,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.LEFT)
	

	Out_Format2(p2_4_2_add1, style,WD_ALIGN_PARAGRAPH.LEFT,0.3)
	Out_Fig1(figures[10], 4,fig8_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)
	
	#Out_Format1(table9_tit,u'黑体' ,False,7,WD_ALIGN_PARAGRAPH.CENTER)
	#Out_table(5,3,table9)

	Out_Fig1(figures[11], 4,fig9_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)
	#Out_Fig1(fig8, 4,fig8_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)
	#Out_Fig1(fig9, 4,fig9_anno,u'黑体',7,WD_ALIGN_PARAGRAPH.CENTER,0.3)

	document.add_heading(h3,1)
	document.add_heading(h3_1,2)
	Out_Format3(p3_1, WD_ALIGN_PARAGRAPH.LEFT)
	document.add_heading(h3_2,2)
	Out_Format3(p3_2,WD_ALIGN_PARAGRAPH.LEFT)

	document.add_heading(h4,1)
	Out_Format3(p4, WD_ALIGN_PARAGRAPH.LEFT)
	
	document.add_heading(h5,2)
	Out_Format3(p5,WD_ALIGN_PARAGRAPH.LEFT)


	document.add_page_break()
	outfile=outpath4proj+'/'+name+'/Report4WESBasic/'+name+'.Report.docx'
	document.save(outfile)
